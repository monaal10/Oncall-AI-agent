"""Word document runbook parser implementation."""

import os
import asyncio
from typing import Dict, List, Optional, Any
from datetime import datetime

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from ..base.runbook_provider import RunbookProvider, RunbookType


class DocxRunbookProvider(RunbookProvider):
    """Word document runbook provider implementation.
    
    Provides integration with DOCX files for runbook content extraction
    and search functionality.
    """

    def __init__(self, config: Dict[str, Any]):
        """Initialize DOCX runbook provider.
        
        Args:
            config: Configuration dictionary containing:
                - runbook_directory: Directory containing DOCX files (required)
                - recursive: Whether to search subdirectories (optional, default: True)
                - cache_enabled: Whether to cache extracted text (optional, default: True)
        """
        if not DOCX_AVAILABLE:
            raise ImportError(
                "DOCX support not available. Install with: pip install python-docx"
            )
        
        super().__init__(config)
        self._text_cache = {} if config.get("cache_enabled", True) else None

    def _validate_config(self) -> None:
        """Validate DOCX runbook configuration.
        
        Raises:
            ValueError: If required configuration is missing
        """
        if "runbook_directory" not in self.config:
            raise ValueError("runbook_directory is required for DOCX runbook provider")
        
        directory = self.config["runbook_directory"]
        if not os.path.exists(directory):
            raise ValueError(f"Runbook directory does not exist: {directory}")
        
        if not os.path.isdir(directory):
            raise ValueError(f"Runbook directory is not a directory: {directory}")

    async def get_runbook_text(
        self,
        runbook_id: str,
        **kwargs
    ) -> str:
        """Get the full text content of a DOCX runbook.
        
        Args:
            runbook_id: DOCX file path (relative to runbook_directory or absolute)
            **kwargs: Additional parameters:
                - include_headers: Whether to include headers/footers (default: False)
                - include_tables: Whether to include table content (default: True)
                
        Returns:
            Full text content of the DOCX file
            
        Raises:
            ConnectionError: If unable to read the DOCX file
            ValueError: If DOCX file not found
        """
        try:
            # Resolve file path
            file_path = self._resolve_file_path(runbook_id)
            
            # Check cache first
            cache_key = f"{file_path}_{hash(str(sorted(kwargs.items())))}"
            if self._text_cache is not None and cache_key in self._text_cache:
                cache_entry = self._text_cache[cache_key]
                # Check if file has been modified since caching
                if cache_entry['modified'] >= os.path.getmtime(file_path):
                    return cache_entry['text']
            
            # Extract text from DOCX
            text = await asyncio.to_thread(
                self._extract_docx_text,
                file_path,
                kwargs.get('include_headers', False),
                kwargs.get('include_tables', True)
            )
            
            # Cache the result
            if self._text_cache is not None:
                self._text_cache[cache_key] = {
                    'text': text,
                    'modified': os.path.getmtime(file_path)
                }
            
            return text
            
        except FileNotFoundError:
            raise ValueError(f"DOCX runbook not found: {runbook_id}")
        except Exception as e:
            raise ConnectionError(f"Failed to read DOCX runbook: {e}")

    async def search_runbooks(
        self,
        query: str,
        limit: Optional[int] = 10
    ) -> List[Dict[str, Any]]:
        """Search for DOCX runbooks containing specific content.
        
        Args:
            query: Search query string
            limit: Maximum number of results to return
            
        Returns:
            List of DOCX runbook search results
            
        Raises:
            ConnectionError: If unable to search runbooks
        """
        try:
            docx_files = await self._find_docx_files()
            results = []
            query_lower = query.lower()
            
            for docx_file in docx_files:
                try:
                    # Get text content
                    text = await self.get_runbook_text(docx_file)
                    text_lower = text.lower()
                    
                    # Check if query matches
                    if query_lower in text_lower:
                        # Extract excerpt around the match
                        excerpt = self._extract_excerpt(text, query, 200)
                        
                        # Calculate relevance score
                        relevance_score = self._calculate_docx_relevance(text_lower, query_lower)
                        
                        # Get file info
                        full_path = self._resolve_file_path(docx_file)
                        file_stat = os.stat(full_path)
                        
                        # Get document metadata
                        metadata = await asyncio.to_thread(self._get_docx_metadata, full_path)
                        
                        result = {
                            'id': docx_file,
                            'title': metadata.get('title') or self._get_docx_title(docx_file),
                            'type': RunbookType.DOCX.value,
                            'excerpt': excerpt,
                            'relevance_score': relevance_score,
                            'last_modified': datetime.fromtimestamp(file_stat.st_mtime),
                            'url': f"file://{full_path}",
                            'file_size': file_stat.st_size,
                            'author': metadata.get('author', 'Unknown')
                        }
                        results.append(result)
                        
                except Exception:
                    # Skip files that can't be processed
                    continue
            
            # Sort by relevance score
            results.sort(key=lambda x: x['relevance_score'], reverse=True)
            
            return results[:limit] if limit else results
            
        except Exception as e:
            raise ConnectionError(f"Failed to search DOCX runbooks: {e}")

    async def list_runbooks(self) -> List[Dict[str, Any]]:
        """List all available DOCX runbooks.
        
        Returns:
            List of DOCX runbooks with metadata
            
        Raises:
            ConnectionError: If unable to list runbooks
        """
        try:
            docx_files = await self._find_docx_files()
            runbooks = []
            
            for docx_file in docx_files:
                try:
                    full_path = self._resolve_file_path(docx_file)
                    file_stat = os.stat(full_path)
                    
                    # Get document metadata
                    metadata = await asyncio.to_thread(self._get_docx_metadata, full_path)
                    
                    runbook = {
                        'id': docx_file,
                        'title': metadata.get('title') or self._get_docx_title(docx_file),
                        'type': RunbookType.DOCX.value,
                        'description': metadata.get('subject', 'Word document runbook'),
                        'last_modified': datetime.fromtimestamp(file_stat.st_mtime),
                        'size': file_stat.st_size,
                        'url': f"file://{full_path}",
                        'author': metadata.get('author', 'Unknown'),
                        'paragraphs': metadata.get('paragraphs', 0),
                        'tables': metadata.get('tables', 0)
                    }
                    runbooks.append(runbook)
                    
                except Exception:
                    # Skip files that can't be processed
                    continue
            
            return runbooks
            
        except Exception as e:
            raise ConnectionError(f"Failed to list DOCX runbooks: {e}")

    async def get_runbook_sections(
        self,
        runbook_id: str,
        **kwargs
    ) -> List[Dict[str, Any]]:
        """Get structured sections from a DOCX runbook.
        
        Args:
            runbook_id: DOCX file identifier
            **kwargs: Additional parameters
            
        Returns:
            List of sections with DOCX-specific metadata
            
        Raises:
            ConnectionError: If unable to access runbook
        """
        try:
            file_path = self._resolve_file_path(runbook_id)
            sections = await asyncio.to_thread(self._extract_docx_sections, file_path)
            return sections
            
        except Exception as e:
            raise ConnectionError(f"Failed to get runbook sections: {e}")

    def _resolve_file_path(self, runbook_id: str) -> str:
        """Resolve runbook ID to full file path.
        
        Args:
            runbook_id: Runbook identifier (file path)
            
        Returns:
            Full path to the DOCX file
        """
        if os.path.isabs(runbook_id):
            return runbook_id
        else:
            return os.path.join(self.config["runbook_directory"], runbook_id)

    async def _find_docx_files(self) -> List[str]:
        """Find all DOCX files in the runbook directory.
        
        Returns:
            List of DOCX file paths (relative to runbook_directory)
        """
        docx_files = []
        directory = self.config["runbook_directory"]
        recursive = self.config.get("recursive", True)
        
        if recursive:
            for root, dirs, files in os.walk(directory):
                for file in files:
                    if file.lower().endswith('.docx') and not file.startswith('~$'):
                        full_path = os.path.join(root, file)
                        relative_path = os.path.relpath(full_path, directory)
                        docx_files.append(relative_path)
        else:
            for file in os.listdir(directory):
                if file.lower().endswith('.docx') and not file.startswith('~$'):
                    full_path = os.path.join(directory, file)
                    if os.path.isfile(full_path):
                        docx_files.append(file)
        
        return docx_files

    def _extract_docx_text(
        self,
        file_path: str,
        include_headers: bool = False,
        include_tables: bool = True
    ) -> str:
        """Extract text from DOCX file.
        
        Args:
            file_path: Path to DOCX file
            include_headers: Whether to include headers/footers
            include_tables: Whether to include table content
            
        Returns:
            Extracted text content
        """
        try:
            doc = Document(file_path)
            text_parts = []
            
            # Extract paragraph text
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract table text if requested
            if include_tables:
                for table in doc.tables:
                    table_text = []
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            cell_text = cell.text.strip()
                            if cell_text:
                                row_text.append(cell_text)
                        if row_text:
                            table_text.append(" | ".join(row_text))
                    
                    if table_text:
                        text_parts.append("\n--- Table ---")
                        text_parts.extend(table_text)
                        text_parts.append("--- End Table ---\n")
            
            # Extract headers/footers if requested
            if include_headers:
                for section in doc.sections:
                    if section.header:
                        for paragraph in section.header.paragraphs:
                            if paragraph.text.strip():
                                text_parts.append(f"[Header] {paragraph.text}")
                    
                    if section.footer:
                        for paragraph in section.footer.paragraphs:
                            if paragraph.text.strip():
                                text_parts.append(f"[Footer] {paragraph.text}")
            
            return '\n'.join(text_parts)
            
        except Exception as e:
            raise ConnectionError(f"Failed to extract text from DOCX: {e}")

    def _extract_docx_sections(self, file_path: str) -> List[Dict[str, Any]]:
        """Extract structured sections from DOCX file.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            List of sections with metadata
        """
        try:
            doc = Document(file_path)
            sections = []
            current_section = None
            paragraph_index = 0
            
            for paragraph in doc.paragraphs:
                paragraph_text = paragraph.text.strip()
                if not paragraph_text:
                    continue
                
                # Check if this is a heading
                is_heading = False
                heading_level = 1
                
                if paragraph.style.name.startswith('Heading'):
                    is_heading = True
                    try:
                        heading_level = int(paragraph.style.name.split()[-1])
                    except (IndexError, ValueError):
                        heading_level = 1
                
                if is_heading:
                    # Save previous section
                    if current_section:
                        sections.append(current_section)
                    
                    # Start new section
                    current_section = {
                        'title': paragraph_text,
                        'level': heading_level,
                        'content': paragraph_text + '\n',
                        'paragraph_start': paragraph_index,
                        'paragraph_end': paragraph_index,
                        'style': paragraph.style.name,
                        'formatting': self._get_paragraph_formatting(paragraph)
                    }
                else:
                    # Add to current section or create default section
                    if current_section:
                        current_section['content'] += paragraph_text + '\n'
                        current_section['paragraph_end'] = paragraph_index
                    else:
                        # Create default section for content without headings
                        current_section = {
                            'title': 'Content',
                            'level': 1,
                            'content': paragraph_text + '\n',
                            'paragraph_start': paragraph_index,
                            'paragraph_end': paragraph_index,
                            'style': paragraph.style.name,
                            'formatting': self._get_paragraph_formatting(paragraph)
                        }
                
                paragraph_index += 1
            
            # Add final section
            if current_section:
                sections.append(current_section)
            
            return sections
            
        except Exception as e:
            raise ConnectionError(f"Failed to extract sections from DOCX: {e}")

    def _get_docx_metadata(self, file_path: str) -> Dict[str, Any]:
        """Get DOCX document metadata.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Dictionary containing document metadata
        """
        metadata = {}
        
        try:
            doc = Document(file_path)
            
            # Core properties
            core_props = doc.core_properties
            metadata.update({
                'title': core_props.title,
                'author': core_props.author,
                'subject': core_props.subject,
                'keywords': core_props.keywords,
                'comments': core_props.comments,
                'category': core_props.category,
                'created': core_props.created,
                'modified': core_props.modified,
                'last_modified_by': core_props.last_modified_by,
                'revision': core_props.revision
            })
            
            # Document statistics
            metadata.update({
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables),
                'sections': len(doc.sections)
            })
            
            # Count non-empty paragraphs
            non_empty_paragraphs = sum(1 for p in doc.paragraphs if p.text.strip())
            metadata['non_empty_paragraphs'] = non_empty_paragraphs
            
        except Exception:
            # If metadata extraction fails, return basic info
            pass
        
        return metadata

    def _get_paragraph_formatting(self, paragraph) -> Dict[str, Any]:
        """Get formatting information for a paragraph.
        
        Args:
            paragraph: DOCX paragraph object
            
        Returns:
            Dictionary containing formatting information
        """
        formatting = {
            'style': paragraph.style.name,
            'alignment': str(paragraph.alignment) if paragraph.alignment else None
        }
        
        # Check for formatting in runs
        if paragraph.runs:
            first_run = paragraph.runs[0]
            formatting.update({
                'bold': first_run.bold,
                'italic': first_run.italic,
                'underline': first_run.underline,
                'font_name': first_run.font.name,
                'font_size': first_run.font.size.pt if first_run.font.size else None
            })
        
        return formatting

    def _get_docx_title(self, docx_file: str) -> str:
        """Generate a title from DOCX filename.
        
        Args:
            docx_file: DOCX file path
            
        Returns:
            Human-readable title
        """
        # Get filename without extension
        filename = os.path.splitext(os.path.basename(docx_file))[0]
        
        # Replace underscores and hyphens with spaces
        title = filename.replace('_', ' ').replace('-', ' ')
        
        # Capitalize words
        title = ' '.join(word.capitalize() for word in title.split())
        
        return title

    def _extract_excerpt(self, text: str, query: str, max_length: int = 200) -> str:
        """Extract an excerpt around the query match.
        
        Args:
            text: Full text content
            query: Search query
            max_length: Maximum excerpt length
            
        Returns:
            Text excerpt containing the query
        """
        text_lower = text.lower()
        query_lower = query.lower()
        
        # Find the first occurrence of the query
        match_pos = text_lower.find(query_lower)
        if match_pos == -1:
            return text[:max_length] + "..." if len(text) > max_length else text
        
        # Calculate excerpt boundaries
        start = max(0, match_pos - max_length // 2)
        end = min(len(text), match_pos + len(query) + max_length // 2)
        
        excerpt = text[start:end]
        
        # Add ellipsis if truncated
        if start > 0:
            excerpt = "..." + excerpt
        if end < len(text):
            excerpt = excerpt + "..."
        
        return excerpt

    def _calculate_docx_relevance(self, text: str, query: str) -> float:
        """Calculate relevance score for DOCX content.
        
        Args:
            text: DOCX text content (lowercase)
            query: Search query (lowercase)
            
        Returns:
            Relevance score between 0 and 1
        """
        if not query:
            return 0.0
        
        # Count occurrences
        query_count = text.count(query)
        if query_count == 0:
            return 0.0
        
        # Base score from occurrence frequency
        text_length = len(text.split())
        frequency_score = min(query_count / max(text_length / 100, 1), 1.0)
        
        # Boost score for matches in headers (indicated by [Header] prefix)
        header_matches = text.count(f"[header] {query}")
        header_boost = min(header_matches * 0.3, 0.5)
        
        # Boost score for matches in tables (indicated by table markers)
        table_context = text[max(0, text.find(query) - 100):text.find(query) + 100]
        table_boost = 0.2 if "--- table ---" in table_context.lower() else 0
        
        return min(frequency_score + header_boost + table_boost, 1.0)
